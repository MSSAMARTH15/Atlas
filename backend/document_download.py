from io import BytesIO

from fastapi import APIRouter
from fastapi.responses import Response
from pydantic import BaseModel
from docx import Document

router = APIRouter()


class DocumentationRequest(BaseModel):
    file_name: str
    documentation: str


@router.post("/download-docs")
def download_docs(request: DocumentationRequest):

    document = Document()

    document.add_heading(
        f"Documentation - {request.file_name}",
        level=1
    )

    document.add_paragraph(request.documentation)

    buffer = BytesIO()
    document.save(buffer)

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{request.file_name}_documentation.docx"'
        }
    )